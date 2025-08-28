# test_word_creator.py
from docx import Document
from plx_parser import parse_plx

def create_word_from_template(template_path, output_path, data):
    """
    Улучшенная функция замены плейсхолдеров
    """
    try:
        doc = Document(template_path)
        
        # Сначала выведем все плейсхолдеры, которые мы ищем
        print("Ищем плейсхолдеры в данных:", list(data.keys()))
        
        # Проходим по всем параграфам документа
        for paragraph in doc.paragraphs:
            original_text = paragraph.text
            for key, value in data.items():
                search_text = "{" + key + "}"
                if search_text in paragraph.text:
                    print(f"Найден плейсхолдер: {search_text} -> {value}")
                    paragraph.text = paragraph.text.replace(search_text, str(value))
            
            if paragraph.text != original_text:
                print(f"Изменен параграф: {original_text} -> {paragraph.text}")
        
        # Проверяем таблицы
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        original_text = paragraph.text
                        for key, value in data.items():
                            search_text = "{" + key + "}"
                            if search_text in paragraph.text:
                                print(f"Найден плейсхолдер в таблице: {search_text} -> {value}")
                                paragraph.text = paragraph.text.replace(search_text, str(value))
                        
                        if paragraph.text != original_text:
                            print(f"Изменена ячейка таблицы: {original_text} -> {paragraph.text}")
        
        doc.save(output_path)
        print(f"Документ сохранен: {output_path}")
        return True
        
    except Exception as e:
        print(f"Ошибка при создании документа: {e}")
        return False

# Тестируем
if __name__ == "__main__":
    # Парсим данные
    all_data = parse_plx("21.03.03_Геодезия_2025_оч_.plx")
    
    if not all_data:
        print("Не удалось распарсить PLX файл")
    else:
        print("Успешно распарсено. Создаем документ...")
        
        # Проверяем, есть ли нужные данные
        if 'ООП' in all_data and all_data['ООП']:
            main_oop_data = all_data['ООП'][0].copy()  # копируем, чтобы не изменять оригинал
            
            print("Данные ООП:", main_oop_data)
            
            # Добавляем компетенции
            competencies = []
            if 'ПланыКомпетенции' in all_data:
                for comp in all_data['ПланыКомпетенции']:
                    comp_text = f"{comp.get('ШифрКомпетенции', '')}: {comp.get('Наименование', '')}"
                    competencies.append(comp_text)
            
            main_oop_data['Компетенции'] = "\n".join(competencies) if competencies else "Не найдены"
            
            # Проверяем, есть ли все необходимые поля
            required_fields = ['Шифр', 'Название', 'УровеньОбразования', 'Квалификация', 
                              'СрокЛет', 'ДатаДокумента', 'Код', 'НомерДокумента', 'ТипГОСа']
            
            print("\nПроверка наличия полей:")
            for field in required_fields:
                if field in main_oop_data:
                    print(f"✓ {field}: {main_oop_data[field]}")
                else:
                    print(f"✗ {field}: ОТСУТСТВУЕТ")
            
            # Создаем документ
            success = create_word_from_template("test.docx", "учебный_план.docx", main_oop_data)
            
            if success:
                print("Документ успешно создан!")
            else:
                print("Не удалось создать документ")
        else:
            print("Не найдены данные ООП в PLX файле")